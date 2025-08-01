# --- Advanced Analytics Stubs ---
def run_predictive(data, question):
    # Predictive analytics: Linear regression (as a starting point)
    import pandas as pd
    import statsmodels.api as sm
    from docx import Document
    import os
    # Assume 'data' is a DataFrame
    df = data if isinstance(data, pd.DataFrame) else pd.DataFrame(data)
    numeric_cols = df.select_dtypes(include='number').columns
    if len(numeric_cols) < 2:
        raise ValueError("Need at least two numeric columns for predictive analytics (regression).")
    y = numeric_cols[0]
    X = numeric_cols[1:]
    X_ = sm.add_constant(df[X])
    model = sm.OLS(df[y], X_, missing='drop').fit()
    summary = model.summary().as_text()
    # Generate report
    doc = Document()
    doc.add_heading('Predictive Analytics Report', 0)
    doc.add_paragraph(f'Question: {question}')
    doc.add_heading('Regression Model', level=1)
    doc.add_paragraph(f'Dependent variable: {y}')
    doc.add_paragraph(f'Independent variables: {", ".join(X)}')
    doc.add_heading('Regression Summary', level=2)
    doc.add_paragraph(summary)
    # Save report with unique filename
    import uuid
    report_path = os.path.join('..', 'reports', f'predictive_report_{uuid.uuid4().hex}.docx')
    doc.save(report_path)
    return report_path

def run_quality_control(data, question):
    # TODO: Implement quality control analytics (SPC, Six Sigma, root cause)
    import uuid, os
    report_path = os.path.join('..', 'reports', f'quality_control_report_{uuid.uuid4().hex}.docx')
    return report_path

def run_machine_learning(data, question):
    # TODO: Implement ML analytics (classification, clustering, anomaly detection)
    import uuid, os
    report_path = os.path.join('..', 'reports', f'ml_report_{uuid.uuid4().hex}.docx')
    return report_path

def run_inventory(data, question):
    # TODO: Implement inventory/supply chain analytics
    import uuid, os
    report_path = os.path.join('..', 'reports', f'inventory_report_{uuid.uuid4().hex}.docx')
    return report_path

def run_workforce(data, question):
    # TODO: Implement workforce/operations analytics
    import uuid, os
    report_path = os.path.join('..', 'reports', f'workforce_report_{uuid.uuid4().hex}.docx')
    return report_path

def run_cost(data, question):
    # TODO: Implement cost/efficiency analytics
    import uuid, os
    report_path = os.path.join('..', 'reports', f'cost_report_{uuid.uuid4().hex}.docx')
    return report_path

def run_rd(data, question):
    # TODO: Implement R&D/product development analytics
    import uuid, os
    report_path = os.path.join('..', 'reports', f'rd_report_{uuid.uuid4().hex}.docx')
    return report_path

def run_environment(data, question):
    # TODO: Implement environmental/compliance analytics
    import uuid, os
    report_path = os.path.join('..', 'reports', f'environment_report_{uuid.uuid4().hex}.docx')
    return report_path

import pandas as pd
import docx
import io
import openai
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import seaborn as sns
import os
import numpy as np
from scipy import stats
import statsmodels.api as sm

# --- File Parsing ---
def process_file(filename, content):
    if filename.endswith('.csv'):
        import tempfile
        with tempfile.NamedTemporaryFile(delete=False, suffix='.csv') as tmp:
            tmp.write(content)
            tmp.flush()
            try:
                df = pd.read_csv(tmp.name, encoding='utf-8', engine='python', on_bad_lines='skip')
                return df
            except Exception:
                try:
                    df = pd.read_csv(tmp.name, encoding='latin1', engine='python', on_bad_lines='skip')
                    return df
                except Exception as e:
                    raise ValueError(f'Could not decode CSV file: {e}')
    elif filename.endswith('.xlsx'):
        df = pd.read_excel(io.BytesIO(content))
        return df
    elif filename.endswith('.docx'):
        doc = docx.Document(io.BytesIO(content))
        text = '\n'.join([p.text for p in doc.paragraphs])
        return text
    elif filename.endswith('.json'):
        import json
        df = pd.read_json(io.BytesIO(content))
        return df
    else:
        # Try to decode as text, else as JSON, else as DataFrame
        try:
            return content.decode(errors='ignore')
        except Exception:
            import json
            try:
                return pd.read_json(io.BytesIO(content))
            except Exception:
                try:
                    return pd.DataFrame(json.loads(content))
                except Exception:
                    return content

# --- Statistical Analysis ---
def analyze_data(df):
    results = {}
    # Descriptive statistics
    results['describe'] = df.describe(include='all').to_dict()
    # Correlation matrix
    results['correlation'] = df.corr(numeric_only=True).to_dict()
    # T-test (if two columns)
    if len(df.select_dtypes(include=np.number).columns) >= 2:
        cols = df.select_dtypes(include=np.number).columns[:2]
        ttest_result = stats.ttest_ind(df[cols[0]].dropna(), df[cols[1]].dropna())
        t_stat = ttest_result[0]
        p_val = ttest_result[1]
        # Ensure t_stat and p_val are always floats
        try:
            t_stat = float(t_stat)
        except Exception:
            t_stat = None
        try:
            p_val = float(p_val)
        except Exception:
            p_val = None
        results['t_test'] = {'columns': list(cols), 't_stat': t_stat, 'p_value': p_val}
    # Regression (if >1 numeric col)
    if len(df.select_dtypes(include=np.number).columns) > 1:
        y = df.select_dtypes(include=np.number).columns[0]
        X = df.select_dtypes(include=np.number).columns[1:]
        X_ = sm.add_constant(df[X])
        model = sm.OLS(df[y], X_, missing='drop').fit()
        results['regression'] = model.summary().as_text()
    return results

# --- NLP Summarization ---
def summarize(text, question):
    prompt = f"Summarize the following data and answer the question: {question}\n\n{text[:2000]}"
    response = openai.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[{"role": "user", "content": prompt}]
    )
    return response.choices[0].message.content

# --- Report Generation ---
def generate_report(summary, question, stats=None, df=None):
    doc = Document()
    doc.add_heading('Auto Statistical Report', 0)
    doc.add_paragraph('Question: ' + question)
    doc.add_heading('Summary', level=1)
    doc.add_paragraph(summary)
    if stats:
        doc.add_heading('Descriptive Statistics', level=2)
        doc.add_paragraph(str(stats.get('describe', {})))
        doc.add_heading('Correlation Matrix', level=2)
        doc.add_paragraph(str(stats.get('correlation', {})))
        if 't_test' in stats:
            doc.add_heading('T-Test', level=2)
            doc.add_paragraph(str(stats['t_test']))
        if 'regression' in stats:
            doc.add_heading('Regression Analysis', level=2)
            doc.add_paragraph(stats['regression'])
    # Add chart if possible
    if df is not None and len(df.select_dtypes(include=np.number).columns) >= 2:
        plt.figure(figsize=(6,4))
        sns.heatmap(df.corr(numeric_only=True), annot=True, cmap='coolwarm')
        plt.title('Correlation Heatmap')
        chart_path = os.path.join('..', 'reports', 'chart.png')
        plt.savefig(chart_path)
        doc.add_picture(chart_path, width=Inches(5))
    # Save report with unique filename
    import uuid
    report_path = os.path.join('..', 'reports', f'report_{uuid.uuid4().hex}.docx')
    doc.save(report_path)
    return report_path
